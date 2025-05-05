from pathlib import Path
import pytesseract
from werkzeug.datastructures import FileStorage
import pdfminer.high_level
from pdfminer.pdfpage import PDFPage
import tempfile
from PIL import Image

from pdf2image import convert_from_path
import docx
from openpyxl import load_workbook


_OCR_LANG = "eng"
_SCANNED_PDF_MIN_CHARS = 50


# Get extract from image using OCR.
def _ocr_image(img):
    return pytesseract.image_to_string(img, lang=_OCR_LANG).strip()


# Extract machine-readable text from PDF.
def _extract_pdf_text(path):
    return pdfminer.high_level.extract_text(str(path)) or ""


# Rudimentary logic to get number of pages to perform OCR on for document scans (minimise processing time).
def pages_to_ocr(total_pages):
    if total_pages == 1:
        return 1
    elif total_pages <= 5:
        return min(2, total_pages)
    else:
        return min(3, total_pages)


# Convert pdf into images.
def _ocr_pdf(path):
    with open(path, "rb") as file:
        num_pages = sum(1 for _ in PDFPage.get_pages(file))

    last_page = pages_to_ocr(num_pages)
    pages = convert_from_path(str(path), dpi=300, first_page=1, last_page=last_page)

    return " ".join(_ocr_image(p) for p in pages)


# Extract text from pdf — fall back on OCR if machine-readable text is sparse.
def _extract_pdf(path):
    text = _extract_pdf_text(path)

    if len(text.strip()) >= _SCANNED_PDF_MIN_CHARS:
        return text

    return _ocr_pdf(path)


# Extract text from docx file paragraphs and table cells.
def _extract_docx(path):
    doc = docx.Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = []

    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                tables.append(row_text)

    return "\n".join(paragraphs + tables)


# Extract text from xlsx file cells.
def _extract_xlsx(path):
    wb = load_workbook(str(path), data_only=True)
    text_lines = []

    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            row_text = " ".join(
                str(cell).strip()
                for cell in row
                if cell is not None and str(cell).strip()
            )
            if row_text:
                text_lines.append(row_text)

    return "\n".join(text_lines)


# Extract text from document, choosing extraction method based on file extension — this is safe as we've already verified the extension matches the MIME type.
def extract_file_text(file, ext):
    ext = ext.lower()

    # Temporarily write file to disk, ensuring all file data is written before performing reads.
    with tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=True) as temp_file:
        file.save(temp_file)
        temp_file.flush()

        path = Path(temp_file.name)

        match ext:
            case "pdf":
                return _extract_pdf(path)
            case "docx":
                return _extract_docx(path)
            case "xlsx":
                return _extract_xlsx(path)
            case "png" | "jpg" | "jpeg" | "tiff" | "bmp":
                return _ocr_image(Image.open(path))
            case _:
                return None
